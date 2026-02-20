from __future__ import annotations

from pathlib import Path

from docx import Document

from ..storage.sqlite import SQLiteStore


def export_comparison_docx(
    store: SQLiteStore,
    dict_ids: list[str],
    out_path: Path,
    mode: str = "table",
) -> Path:
    rows = store.comparison_rows(dict_ids)

    doc = Document()
    doc.add_heading("DicoDiachro Multi-comparison", level=1)
    doc.add_paragraph(f"Dictionaries: {', '.join(dict_ids)}")

    if mode == "list":
        for row in rows:
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"{row['lemma_label']} ({row['lemma_group_id'][:8]})")
            for dict_id in dict_ids:
                value = row["values"].get(dict_id, "ABSENT")
                doc.add_paragraph(f"{dict_id}: {value}", style="List Bullet")
    else:
        table = doc.add_table(rows=1, cols=2 + len(dict_ids))
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "lemma_group_id"
        headers[1].text = "lemma_label"
        for idx, dict_id in enumerate(dict_ids, start=2):
            headers[idx].text = dict_id

        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row["lemma_group_id"]
            cells[1].text = row["lemma_label"]
            for idx, dict_id in enumerate(dict_ids, start=2):
                cells[idx].text = row["values"].get(dict_id, "ABSENT")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
